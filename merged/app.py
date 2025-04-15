import os
import time
import pandas as pd
import tempfile
import docx
import json
import requests
from datetime import datetime
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options as ChromeOptions
from selenium.webdriver.firefox.options import Options as FirefoxOptions
from selenium.webdriver.edge.options import Options as EdgeOptions
from webdriver_manager.chrome import ChromeDriverManager
from webdriver_manager.firefox import GeckoDriverManager
from webdriver_manager.microsoft import EdgeChromiumDriverManager
from flask import Flask, render_template, request, redirect, url_for, send_file, flash,jsonify
from docx import Document
import google.generativeai as genai
from dotenv import load_dotenv
import subprocess




# Load environment variables
load_dotenv(verbose=False)

# Initialize Flask app
app = Flask(__name__)
app.secret_key = "testing_automation_secret_key"

@app.route("/")
def home():
    return render_template("index.html")  # Landing page

@app.route("/run_test", methods=["POST"])
def run_test():
    data = request.get_json()
    url = data["url"]
    username = data["username"]
    password = data["password"]

    try:
        result = subprocess.run(["your_command_here"], capture_output=True, text=True, check=True)
        return "Test Success"  # Explicitly return success only if no error occurs
    except subprocess.CalledProcessError as e:
         return f"Test Failed: {e.stderr}"  # Indicate failure properly
    
# Define file paths and configurations
test_cases_file = "test_cases.csv"
UPLOAD_FOLDER = os.path.join(os.getcwd(), 'generated_files')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER



def load_data():
    """Load and process test case data from CSV"""
    if not os.path.exists(test_cases_file):
        return {"total": 0, "passed": 0, "failed": 0, "details": []}

    df = pd.read_csv(test_cases_file)

    # Convert actual_result to lowercase for consistency
    df["actual_result"] = df["actual_result"].str.lower()

    total_tests = len(df)
    passed_tests = len(df[df["actual_result"] == "success"])
    failed_tests = len(df[df["actual_result"] == "failed"])

    details = df[["id", "description", "actual_result"]].to_dict(orient="records")

    return {
        "total": total_tests,
        "passed": passed_tests,
        "failed": failed_tests,
        "details": details
    }




@app.route('/get_chart_data')
def get_chart_data():
    # Simulated dynamic test data
    test_data = {
        'doughnut': [70, 30],  # Passed, Failed
        'gauge': 70,  # Percentage of passed tests
        'area': {
            'dates': ["Mar 1", "Mar 5", "Mar 10", "Mar 15", "Mar 20", "Mar 25", "Mar 30"],
            'tests_executed': [10, 25, 35, 50, 65, 80, 100]
        }
    }
    return jsonify(test_data)  # Send JSON data to frontend

# --------------------- TEST AUTOMATION FUNCTIONS ---------------------

# ✅ Generate Unique Report File
def get_report_filename():
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"test_report_{timestamp}.docx"


# ✅ Setup Selenium WebDriver with Cross-Browser Support and Custom URL
def setup_driver(website_url=None, browser_type="chrome", headless=False):
    driver = None
    
    try:
        # Setup Chrome browser
        if browser_type.lower() == "chrome":
            chrome_options = ChromeOptions()
            if headless:
                chrome_options.add_argument("--headless")
            chrome_options.add_argument("--window-size=1920,1080")
            chrome_options.add_argument("--disable-gpu")
            chrome_options.add_argument("--no-sandbox")
            chrome_options.add_argument("--disable-dev-shm-usage")
            driver = webdriver.Chrome(ChromeDriverManager().install(), options=chrome_options)
        
        # Setup Firefox browser
        elif browser_type.lower() == "firefox":
            firefox_options = FirefoxOptions()
            if headless:
                firefox_options.add_argument("--headless")
            driver = webdriver.Firefox(executable_path=GeckoDriverManager().install(), options=firefox_options)
        
        # Setup Edge browser
        elif browser_type.lower() == "edge":
            edge_options = EdgeOptions()
            if headless:
                edge_options.add_argument("--headless")
            driver = webdriver.Edge(EdgeChromiumDriverManager().install(), options=edge_options)
        
        # Default to Chrome if browser type not recognized
        else:
            print(f"⚠️ Browser type '{browser_type}' not recognized. Using Chrome.")
            chrome_options = ChromeOptions()
            if headless:
                chrome_options.add_argument("--headless")
            driver = webdriver.Chrome(ChromeDriverManager().install(), options=chrome_options)
        
        # Use the provided URL or default if none is provided
        url = website_url if website_url else "https://gdgcmarwadiuniversity.tech/admin/login.php"
        driver.get(url)
        
        # Set an implicit wait to handle dynamic elements
        driver.implicitly_wait(10)
        
        return driver
    
    except Exception as e:
        print(f"❌ Error setting up {browser_type} driver: {str(e)}")
        if driver:
            driver.quit()
        # Fall back to basic Chrome setup without webdriver manager if there's an error
        driver = webdriver.Chrome()
        url = website_url if website_url else "https://gdgcmarwadiuniversity.tech/admin/login.php"
        driver.get(url)
        return driver


# ✅ Logout Function (Enhanced)
def logout(driver, logout_url=None):
    try:
        # Try to find and click logout button
        logout_elements = driver.find_elements(By.XPATH, "//a[contains(text(), 'Logout') or contains(text(), 'Log out')]")
        
        if logout_elements:
            logout_elements[0].click()
            time.sleep(1)
            return True

        # If provided a specific logout URL, navigate there
        if logout_url:
            driver.get(logout_url)
            time.sleep(1)
            return True
            
        # If no logout button found and no URL provided, go back to login page
        driver.get(driver.current_url)  # Just refresh the page
        time.sleep(1)
        return True

    except Exception as e:
        print(f"❌ Logout Error: {str(e)}")
        # Just refresh current page
        driver.get(driver.current_url)
        time.sleep(1)
        return False


# ✅ Run Individual Test Case with Enhanced Element Finding
def run_test_case(driver, test_case):
    actual_result = "Failed"  

    try:
        print(f"🟢 Running Test: {test_case['id']} - {test_case['description']}")

        # Get the URL from test case or use current URL
        login_url = test_case.get('website_url', driver.current_url)
        
        # Ensure we're on the correct page
        if driver.current_url != login_url:
            driver.get(login_url)
            time.sleep(2)

        # Get field information
        username_field = test_case.get('username_field', 'username')
        password_field = test_case.get('password_field', 'password')
        submit_button = test_case.get('submit_button', 'login')
        
        # Convert success_indicators to string before splitting (fix for float error)
        success_indicators_value = test_case.get('success_indicators', 'dashboard,home,profile,welcome,index')
        success_indicators = str(success_indicators_value).split(',')
        
        # Enhanced approach to find username input field - try multiple strategies
        username_input = None
        username_find_methods = [
            # Try by name
            lambda: driver.find_element(By.NAME, username_field),
            # Try by ID
            lambda: driver.find_element(By.ID, username_field),
            # Try by CSS selector matching name or id
            lambda: driver.find_element(By.CSS_SELECTOR, f"input[name='{username_field}'], input[id='{username_field}']"),
            # Try by common input types with placeholders
            lambda: driver.find_element(By.XPATH, f"//input[@placeholder='{username_field}']"),
            # Try by common username attributes
            lambda: driver.find_element(By.XPATH, "//input[@type='text' or @type='email' or @name='username' or @id='username' or @name='email' or @id='email']"),
            # Try any input with labels containing username/email
            lambda: driver.find_element(By.XPATH, "//label[contains(translate(text(), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), 'username') or contains(translate(text(), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), 'email')]/following::input[1]")
        ]
        
        for find_method in username_find_methods:
            try:
                username_input = find_method()
                if username_input:
                    break
            except:
                continue
        
        if not username_input:
            raise Exception("Username input field not found")
            
        # Enhanced approach to find password input field - try multiple strategies
        password_input = None
        password_find_methods = [
            # Try by name
            lambda: driver.find_element(By.NAME, password_field),
            # Try by ID
            lambda: driver.find_element(By.ID, password_field),
            # Try by CSS selector matching name or id
            lambda: driver.find_element(By.CSS_SELECTOR, f"input[name='{password_field}'], input[id='{password_field}']"),
            # Try by placeholder
            lambda: driver.find_element(By.XPATH, f"//input[@placeholder='{password_field}']"),
            # Try by type=password
            lambda: driver.find_element(By.XPATH, "//input[@type='password']"),
            # Try any input with labels containing password
            lambda: driver.find_element(By.XPATH, "//label[contains(translate(text(), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), 'password')]/following::input[1]")
        ]
        
        for find_method in password_find_methods:
            try:
                password_input = find_method()
                if password_input:
                    break
            except:
                continue
        
        if not password_input:
            raise Exception("Password input field not found")
        
        # Enhanced approach to find login button - try multiple strategies
        login_button = None
        button_find_methods = [
            # Try by name
            lambda: driver.find_element(By.NAME, submit_button),
            # Try by ID
            lambda: driver.find_element(By.ID, submit_button),
            # Try by button/input text/value
            lambda: driver.find_element(By.XPATH, f"//button[contains(translate(text(), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), '{submit_button.lower()}')]"),
            lambda: driver.find_element(By.XPATH, f"//input[@value='{submit_button}' or contains(translate(@value, 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), '{submit_button.lower()}')]"),
            # Try common submit buttons
            lambda: driver.find_element(By.XPATH, "//button[@type='submit']"),
            lambda: driver.find_element(By.XPATH, "//input[@type='submit']"),
            # Try common login button texts
            lambda: driver.find_element(By.XPATH, "//button[contains(translate(text(), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), 'login') or contains(translate(text(), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), 'log in') or contains(translate(text(), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), 'sign in')]"),
            lambda: driver.find_element(By.XPATH, "//input[@value='Login' or @value='Log In' or @value='Sign In' or @value='Submit']")
        ]
        
        for find_method in button_find_methods:
            try:
                login_button = find_method()
                if login_button:
                    break
            except:
                continue
        
        if not login_button:
            raise Exception("Login button not found")

        username_input.clear()
        password_input.clear()

        username_input.send_keys(str(test_case['username']))
        password_input.send_keys(str(test_case['password']))
        login_button.click()

        time.sleep(3)  # Allow time for page to load after clicking submit

        current_url = driver.current_url.lower()
        page_source = driver.page_source.lower()

        # ✅ Success condition: Redirected to dashboard/home or success indicators are present
        if any(word in current_url for word in success_indicators) or \
           any(word in page_source for word in ['welcome', 'dashboard', 'logged in', 'successfully']):
            actual_result = "Success"

        # ✅ Failure condition: Error messages detected
        if any(word in page_source for word in ["error", "invalid", "incorrect", "failed", "wrong password"]):
            actual_result = "Failed"

        print(f"✅ {test_case['id']} - Expected: {test_case['expected']}, Actual: {actual_result}")

        # Get logout URL from test case or use None
        logout_url = test_case.get('logout_url', None)
        logout(driver, logout_url)

    except Exception as e:
        actual_result = f"Error: {str(e)}"
        print(f"❌ Error in test {test_case['id']}: {str(e)}")
        logout(driver)

    return actual_result


# ✅ Save Results to CSV
def save_results_to_csv(test_cases):
    df = pd.DataFrame(test_cases)
    try:
        df.to_csv(test_cases_file, index=False)
        print(f"✅ Results saved to {test_cases_file}")
    except Exception as e:
        print(f"❌ CSV Save Error: {str(e)}")


# ✅ Generate Report (Tabular Format with Password)
def generate_report(test_cases, report_path, browser_results=None):
    doc = Document()
    doc.add_heading("Test Report", level=1)
    
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_paragraph(f"📅 Report Generated On: {timestamp}\n")

    # Add cross-browser testing results if available
    if browser_results:
        doc.add_heading("Cross-Browser Test Results", level=2)
        browser_table = doc.add_table(rows=1, cols=4)
        browser_table.style = 'Table Grid'
        
        # Add headers
        hdr_cells = browser_table.rows[0].cells
        browser_headers = ["Browser", "Tests Run", "Tests Passed", "Pass Rate"]
        for i, header in enumerate(browser_headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        # Add data for each browser
        for browser, results in browser_results.items():
            row_cells = browser_table.add_row().cells
            row_cells[0].text = browser
            row_cells[1].text = str(results.get("total", 0))
            row_cells[2].text = str(results.get("passed", 0))
            pass_rate = "0%" if results.get("total", 0) == 0 else f"{(results.get('passed', 0) / results.get('total', 0) * 100):.1f}%"
            row_cells[3].text = pass_rate
        
        doc.add_paragraph("")  # Add spacing

    # Add detailed test results
    doc.add_heading("Detailed Test Results", level=2)
    table = doc.add_table(rows=1, cols=8)  # Added columns for browser and website URL
    table.style = 'Table Grid'

    headers = ["Test ID", "Description", "Website URL", "Browser", "Username", "Password", "Expected Result", "Actual Result"]
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    for test in test_cases:
        row_cells = table.add_row().cells
        row_cells[0].text = str(test.get("id", "N/A"))
        row_cells[1].text = str(test.get("description", "N/A"))
        row_cells[2].text = str(test.get("website_url", "Default"))
        row_cells[3].text = str(test.get("browser", "Chrome"))  # New column for browser
        row_cells[4].text = str(test.get("username", "N/A"))
        row_cells[5].text = str(test.get("password", "******"))  # Mask password
        row_cells[6].text = str(test.get("expected", "N/A"))
        row_cells[7].text = str(test.get("actual_result", "N/A"))

    os.makedirs(os.path.dirname(report_path) if os.path.dirname(report_path) else '.', exist_ok=True)
    
    doc.save(report_path)
    print(f"📄 Report saved as {report_path}")


# --------------------- TEST CASE GENERATION FUNCTIONS ---------------------

def save_to_docx(page_name, test_cases, file_path):
    """Save the generated test cases to a Word document in tabular format"""

    # Ensure the directory exists
    os.makedirs(os.path.dirname(file_path), exist_ok=True)

    # Create a new document
    doc = docx.Document()
    
    # Add title
    doc.add_heading(f"Test Cases for {page_name}", 0)
    
    # Add generation timestamp
    doc.add_paragraph(f"Generated on: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()
    
    # Add table
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    
    # Define column headers
    hdr_cells = table.rows[0].cells
    headers = ["Test Case ID", "Description", "Preconditions", "Test Steps", "Expected Results", "Test Data", "Priority"]
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    
    # Split test cases into table rows
    test_cases_lines = test_cases.strip().split('\n')
    for line in test_cases_lines:
        if line.strip():
            row_cells = table.add_row().cells
            case_data = line.split('|')  # Assuming '|' is used as a separator in the response
            for i in range(min(len(case_data), 7)):
                if i < len(case_data):
                    row_cells[i].text = case_data[i].strip()
                else:
                    row_cells[i].text = ""
    
    # Save the document
    doc.save(file_path)


# --------------------- FIGMA FILE PROCESSING FUNCTIONS ---------------------

def get_figma_file_data(file_key, access_token):
    """Fetch Figma file data using the Figma API"""
    headers = {
        'X-Figma-Token': access_token
    }
    response = requests.get(f'https://api.figma.com/v1/files/{file_key}', headers=headers)
    
    if response.status_code != 200:
        raise Exception(f"Failed to fetch Figma file: {response.text}")
    
    return response.json()


def extract_figma_content(figma_data):
    """Extract relevant content from Figma file data"""
    document = figma_data.get('document', {})
    name = figma_data.get('name', 'Figma Design')
    
    # Extract text nodes, frame names, and component names
    extracted_content = {
        'name': name,
        'screens': [],
        'components': [],
        'texts': []
    }
    
    # Recursive function to traverse Figma document
    def traverse_node(node):
        if node.get('type') == 'FRAME' or node.get('type') == 'COMPONENT':
            node_info = {
                'name': node.get('name', ''),
                'type': node.get('type', ''),
                'children_count': len(node.get('children', [])),
                'id': node.get('id', '')
            }
            
            if node.get('type') == 'FRAME':
                extracted_content['screens'].append(node_info)
            else:
                extracted_content['components'].append(node_info)
        
        # Extract text content
        if node.get('type') == 'TEXT':
            extracted_content['texts'].append({
                'content': node.get('characters', ''),
                'style': node.get('style', {})
            })
        
        # Traverse children
        for child in node.get('children', []):
            traverse_node(child)
    
    # Start traversal from the top-level children
    for child in document.get('children', []):
        traverse_node(child)
    
    return extracted_content


def generate_test_cases_from_figma(figma_content, model):
    """Generate test cases based on Figma content using Gemini model"""
    
    # Create a comprehensive prompt based on Figma content
    screens = "\n".join([f"- {screen['name']}" for screen in figma_content['screens']])
    components = "\n".join([f"- {comp['name']}" for comp in figma_content['components']])
    texts = "\n".join([f"- {text['content']}" for text in figma_content['texts'] if text['content'].strip()])
    
    prompt = f"""Generate comprehensive test cases for a UI design titled '{figma_content['name']}' in a tabular format.

Content extracted from the design:

Screens/Frames ({len(figma_content['screens'])}):
{screens}

UI Components ({len(figma_content['components'])}):
{components}

Text Content:
{texts}

Include the following columns:
- Test Case ID
- Test Case Description
- Preconditions
- Test Steps
- Expected Results
- Test Data
- Priority (High/Medium/Low)

Cover functional testing, UI/UX testing, error handling, and edge cases for all screens and interactions.
Make the test cases specific to this design, using the extracted information.
"""
    
    # Generate content with Gemini
    response = model.generate_content(prompt)
    
    return response.text


# --------------------- FLASK ROUTES ---------------------

# ✅ Main Index Page with functionalities
# ✅ Dashboard Page (Old Index)
@app.route('/dashboard')
def dashboard():
    # Ensure test_cases_file exists to avoid FileNotFoundError
    tests = []
    if os.path.exists(test_cases_file):
        tests = pd.read_csv(test_cases_file, dtype=str).to_dict(orient='records')

    # API keys
    api_key = os.getenv("GEMINI_API_KEY", "AIzaSyAHdc9rZS0HN10XYpkOPQTvuWuc6kO4nBM")
    figma_token = os.getenv("FIGMA_ACCESS_TOKEN", "figd_QGsUs3wlfl4kQ3qlD2nhohiIESPLPRYsFBpHLIDq")

    return render_template('dashboard.html', tests=tests, api_key=api_key, figma_token=figma_token)

# ✅ Ensure Static Files Work (CSS, JS, Images)
@app.route('/static/<path:filename>')
def static_files(filename):
    return url_for('static', filename=filename)


# ✅ Add Test Cases Route
@app.route('/add_test_cases', methods=['POST'])
def add_test_cases():
    test_ids = request.form.getlist('test_id[]')
    descriptions = request.form.getlist('description[]')
    website_urls = request.form.getlist('website_url[]')
    usernames = request.form.getlist('username[]')
    passwords = request.form.getlist('password[]')
    username_fields = request.form.getlist('username_field[]')
    password_fields = request.form.getlist('password_field[]')
    submit_buttons = request.form.getlist('submit_button[]')
    success_indicators = request.form.getlist('success_indicators[]')
    logout_urls = request.form.getlist('logout_url[]')
    expecteds = request.form.getlist('expected[]')
    browsers = request.form.getlist('browser[]')  # New field for browser

    new_tests = []
    for i in range(len(test_ids)):
        new_tests.append({
            "id": test_ids[i],
            "description": descriptions[i],
            "website_url": website_urls[i],
            "username": usernames[i],
            "password": passwords[i],
            "username_field": username_fields[i] if username_fields and i < len(username_fields) else "username",
            "password_field": password_fields[i] if password_fields and i < len(password_fields) else "password",
            "submit_button": submit_buttons[i] if submit_buttons and i < len(submit_buttons) else "login",
            "success_indicators": success_indicators[i] if success_indicators and i < len(success_indicators) else "dashboard,home,profile,welcome,index",
            "logout_url": logout_urls[i] if logout_urls and i < len(logout_urls) else "",
            "expected": expecteds[i],
            "browser": browsers[i] if browsers and i < len(browsers) else "chrome"  # Default to Chrome
        })

    os.makedirs(os.path.dirname(test_cases_file) if os.path.dirname(test_cases_file) else '.', exist_ok=True)

    existing_tests = []
    if os.path.exists(test_cases_file):
        existing_tests = pd.read_csv(test_cases_file, dtype=str).to_dict(orient='records')

    all_tests = existing_tests + new_tests

    pd.DataFrame(all_tests).to_csv(test_cases_file, index=False)
    flash("Test cases successfully added!")
    return redirect(url_for('dashboard'))


# ✅ Run Tests Route with Cross-Browser Support
@app.route('/run_tests')
def run_tests():
    if not os.path.exists(test_cases_file):
        flash("No test cases found.")
        return redirect(url_for('dashboard'))

    test_cases = pd.read_csv(test_cases_file, dtype=str).to_dict(orient='records')
    
    # Group test cases by browser
    browser_groups = {}
    for test in test_cases:
        browser = test.get('browser', 'chrome').lower()
        if browser not in browser_groups:
            browser_groups[browser] = []
        browser_groups[browser].append(test)
    
    # Add default group if none specified
    if not browser_groups:
        browser_groups['chrome'] = test_cases
    
    # Track results by browser
    browser_results = {}
    
    # Run tests for each browser type
    for browser_type, browser_tests in browser_groups.items():
        driver = None
        browser_results[browser_type] = {
            "total": len(browser_tests),
            "passed": 0
        }
        
        try:
            previous_url = None
            for test in browser_tests:
                # Create a new driver instance for each website or reuse for the same website
                current_url = test.get('website_url')
                
                if not driver or (current_url and current_url != previous_url):
                    if driver:
                        driver.quit()
                    driver = setup_driver(current_url, browser_type)
                    previous_url = current_url
                
                # Run the test case
                result = run_test_case(driver, test)
                test["actual_result"] = result
                
                # Track passed tests
                if result == "Success":
                    browser_results[browser_type]["passed"] += 1
        
        except Exception as e:
            print(f"❌ Error running tests in {browser_type}: {str(e)}")
        
        finally:
            if driver:
                driver.quit()
    
    # Update all test cases with results
    save_results_to_csv(test_cases)

    report_file = get_report_filename()
    
    try:
        generate_report(test_cases, report_file, browser_results)
        flash("Cross-browser tests completed successfully! Report generated.")
        return redirect(url_for('dashboard', download_report=report_file))
    except Exception as e:
        print(f"Error generating report: {e}")
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_report_path = temp_file.name
        
        generate_report(test_cases, temp_report_path, browser_results)
        flash("Tests completed with errors. Report generated.")
        return redirect(url_for('dashboard', download_report=temp_report_path))


# ✅ Download Report Route
@app.route('/download_report/<path:filename>')
def download_report(filename):
    return send_file(filename, as_attachment=True)


# ✅ Clear All Test Cases Route
@app.route('/clear_tests', methods=['POST'])
def clear_tests():
    try:
        if os.path.exists(test_cases_file):
            os.remove(test_cases_file)
            flash("All test cases have been cleared successfully!")
        else:
            flash("No test cases file exists.")
    except Exception as e:
        flash(f"Error clearing test cases: {str(e)}")
    
    return redirect(url_for('dashboard'))


# ✅ Clear Selected Test Cases Route
@app.route('/clear_selected_tests', methods=['POST'])
def clear_selected_tests():
    if not os.path.exists(test_cases_file):
        flash("No test cases file exists.")
        return redirect(url_for('dashboard'))
    
    selected_ids = request.form.getlist('selected_tests[]')
    
    if not selected_ids:
        flash("No test cases were selected.")
        return redirect(url_for('dashboard'))
    
    try:
        # Read existing test cases
        df = pd.read_csv(test_cases_file, dtype=str)
        
        # Filter out selected test cases
        df = df[~df['id'].isin(selected_ids)]
        
        if df.empty:
            # If all tests were removed, delete the file
            os.remove(test_cases_file)
            flash("All selected test cases have been cleared successfully!")
        else:
            # Save the filtered dataframe back to CSV
            df.to_csv(test_cases_file, index=False)
            flash(f"Selected test cases ({len(selected_ids)}) have been cleared successfully!")
    
    except Exception as e:
        flash(f"Error clearing selected test cases: {str(e)}")
    
    return redirect(url_for('dashboard'))


# ✅ Generate Test Cases Route
@app.route('/generate_test_cases', methods=['POST'])
def generate_test_cases():
    # Get form data
    page_name = request.form.get('page_name', '').strip()
    additional_context = request.form.get('additional_context', '').strip()
    
    # Validate input
    if not page_name:
        flash('Please enter a page name', 'error')
        return redirect(url_for('dashboard'))
    
    try:
        # Configure the Gemini API
        api_key = os.getenv("GEMINI_API_KEY", 'AIzaSyAHdc9rZS0HN10XYpkOPQTvuWuc6kO4nBM')
        genai.configure(api_key=api_key)
        
        # Setup the model (Updated for Gemini)
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        # Build prompt for Gemini
        prompt = f"""Generate comprehensive test cases for a webpage named '{page_name}' in a tabular format.
        
        Include the following columns:
        - Test Case ID
        - Test Case Description
        - Preconditions
        - Test Steps
        - Expected Results
        - Test Data
        - Priority (High/Medium/Low)
        
        Cover functional testing, UI/UX testing, error handling, and edge cases.

Additional context about the webpage:
{additional_context}

Format each test case as a single line with pipe (|) separators between columns.
"""

        # Generate test cases using Gemini
        response = model.generate_content(prompt)
        test_cases = response.text

        # Create filename based on page name
        safe_filename = "".join(c if c.isalnum() else "_" for c in page_name)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{safe_filename}_test_cases.docx")

        # Save to Word document
        save_to_docx(page_name, test_cases, file_path)

        flash(f"Test cases generated successfully for {page_name}!", "success")
        return send_file(file_path, as_attachment=True)
    
    except Exception as e:
        flash(f"Error generating test cases: {str(e)}", "error")
        return redirect(url_for('dashboard'))


# ✅ Generate Test Cases from Figma Route
@app.route('/generate_from_figma', methods=['POST'])
def generate_from_figma():
    figma_file_key = request.form.get('figma_file_key', '').strip()
    figma_access_token = request.form.get('figma_access_token', '').strip()
    
    if not figma_file_key or not figma_access_token:
        flash('Please provide both Figma file key and access token', 'error')
        return redirect(url_for('dashboard'))
    
    try:
        # Configure Gemini
        api_key = os.getenv("GEMINI_API_KEY", 'AIzaSyAHdc9rZS0HN10XYpkOPQTvuWuc6kO4nBM')
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        # Fetch and process Figma file
        figma_data = get_figma_file_data(figma_file_key, figma_access_token)
        figma_content = extract_figma_content(figma_data)
        
        # Generate test cases
        test_cases = generate_test_cases_from_figma(figma_content, model)
        
        # Create file name based on Figma file name
        safe_filename = "".join(c if c.isalnum() else "_" for c in figma_content['name'])
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{safe_filename}_figma_test_cases.docx")
        
        # Save to docx
        save_to_docx(figma_content['name'], test_cases, file_path)
        
        flash(f"Test cases generated successfully from Figma design: {figma_content['name']}!", "success")
        return send_file(file_path, as_attachment=True)
        
    except Exception as e:
        flash(f"Error generating test cases from Figma: {str(e)}", "error")
        return redirect(url_for('dashboard'))


# ✅ Upload Test Cases Route
@app.route('/upload_test_cases', methods=['POST'])
def upload_test_cases():
    # Check if a file was uploaded
    if 'test_cases_file' not in request.files:
        flash("No file uploaded", "error")
        return redirect(url_for('dashboard'))
    
    uploaded_file = request.files['test_cases_file']
    
    # Check if a file was selected
    if uploaded_file.filename == '':
        flash("No file selected", "error")
        return redirect(url_for('dashboard'))
    
    try:
        # Create a temporary file path
        temp_file_path = os.path.join(tempfile.gettempdir(), uploaded_file.filename)
        uploaded_file.save(temp_file_path)
        
        # Process based on file extension
        file_ext = os.path.splitext(uploaded_file.filename)[1].lower()
        
        if file_ext == '.csv':
            # Read CSV directly
            df = pd.read_csv(temp_file_path, dtype=str)
            
            # Validate that it contains required columns
            required_columns = ["id", "description", "username", "password", "expected"]
            missing_columns = [col for col in required_columns if col not in df.columns]
            
            if missing_columns:
                flash(f"CSV is missing required columns: {', '.join(missing_columns)}", "error")
                return redirect(url_for('dashboard'))
            
            # Add default values for optional columns if they're missing
            for column, default in [
                ("website_url", "https://gdgcmarwadiuniversity.tech/admin/login.php"),
                ("username_field", "username"),
                ("password_field", "password"),
                ("submit_button", "login"),
                ("success_indicators", "dashboard,home,profile,welcome,dashboard"),
                ("browser", "chrome")
            ]:
                if column not in df.columns:
                    df[column] = default
        
        elif file_ext in ['.xlsx', '.xls']:
            # Read Excel file
            df = pd.read_excel(temp_file_path, dtype=str)
            
            # Validate that it contains required columns
            required_columns = ["id", "description", "username", "password", "expected"]
            missing_columns = [col for col in required_columns if col not in df.columns]
            
            if missing_columns:
                flash(f"Excel file is missing required columns: {', '.join(missing_columns)}", "error")
                return redirect(url_for('dashboard'))
            
            # Add default values for optional columns if they're missing
            for column, default in [
                ("website_url", "https://gdgcmarwadiuniversity.tech/admin/login.php"),
                ("username_field", "username"),
                ("password_field", "password"),
                ("submit_button", "login"),
                ("success_indicators", "dashboard,home,profile,welcome,dashboard"),
                ("browser", "chrome")
            ]:
                if column not in df.columns:
                    df[column] = default
        
        else:
            flash(f"Unsupported file format: {file_ext}. Please upload a CSV or Excel file.", "error")
            return redirect(url_for('dashboard'))
        
        # Save to test cases file
        if os.path.exists(test_cases_file):
            # Append to existing file
            existing_df = pd.read_csv(test_cases_file, dtype=str)
            combined_df = pd.concat([existing_df, df], ignore_index=True)
            combined_df.to_csv(test_cases_file, dashboard=False)
        else:
            # Create new file
            df.to_csv(test_cases_file,index=False)
        
        flash(f"Successfully uploaded {len(df)} test cases", "success")
        return redirect(url_for('dashboard'))
    
    except Exception as e:
        flash(f"Error processing uploaded file: {str(e)}", "error")
        return redirect(url_for('dashboard'))
    
    finally:
        # Clean up temporary file
        if 'temp_file_path' in locals() and os.path.exists(temp_file_path):
            os.remove(temp_file_path)


# ✅ Download Template Route
@app.route('/download_template')
def download_template():
    # Create a pandas DataFrame with the template structure
    template_df = pd.DataFrame({
        "id": ["TC001", "TC002"],
        "description": ["Valid Login Test", "Invalid Login Test"],
        "website_url": ["https://example.com/login", "https://example.com/login"],
        "username": ["test_user", "invalid_user"],
        "password": ["password123", "wrongpass"],
        "username_field": ["username", "username"],
        "password_field": ["password", "password"],
        "submit_button": ["login", "login"],
        "success_indicators": ["dashboard,home", "dashboard,home"],
        "logout_url": ["", ""],
        "expected": ["Success", "Failed"],
        "browser": ["chrome", "firefox"]
    })
    
    # Create a temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".csv")
    template_path = temp_file.name
    temp_file.close()
    
    # Save DataFrame to the temporary file
    template_df.to_csv(template_path, index=False)
    
    # Send the file to the user
    return send_file(template_path, as_attachment=True, 
                     download_name="test_cases_template.csv",
                     mimetype="text/csv")


# ✅ Export Test Cases Route
@app.route('/export_test_cases')
def export_test_cases():
    if not os.path.exists(test_cases_file):
        flash("No test cases to export", "error")
        return redirect(url_for('dashboard'))
    
    try:
        # Get the format parameter from the query string
        export_format = request.args.get('format', 'csv').lower()
        
        # Create a temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False)
        temp_path = temp_file.name
        temp_file.close()
        
        # Read the test cases
        df = pd.read_csv(test_cases_file, dtype=str)
        
        if export_format == 'excel':
            # Export as Excel
            df.to_excel(temp_path, index=False, engine='openpyxl')
            mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            download_name = 'test_cases_export.xlsx'
        elif export_format == 'json':
            # Export as JSON
            with open(temp_path, 'w') as f:
                df.to_json(f, orient='records', indent=4)
            mimetype = 'application/json'
            download_name = 'test_cases_export.json'
        else:
            # Default: Export as CSV
            df.to_csv(temp_path, index=False)
            mimetype = 'text/csv'
            download_name = 'test_cases_export.csv'
        
        return send_file(temp_path, as_attachment=True, 
                         download_name=download_name,
                         mimetype=mimetype)
    
    except Exception as e:
        flash(f"Error exporting test cases: {str(e)}", "error")
        return redirect(url_for('dashboard'))


# ✅ API Route to run tests and return JSON results
@app.route('/api/run_tests', methods=['POST'])
def api_run_tests():
    # Get JSON data from request
    request_data = request.json or {}
    
    # Get test cases either from the request or from the file
    test_cases = request_data.get('test_cases', [])
    
    if not test_cases:
        if os.path.exists(test_cases_file):
            test_cases = pd.read_csv(test_cases_file, dtype=str).to_dict(orient='records')
        else:
            return jsonify({
                'success': False,
                'message': 'No test cases provided or found',
                'results': []
            }), 400
    
    browser_type = request_data.get('browser', 'chrome').lower()
    headless = request_data.get('headless', True)
    
    # Initialize results
    results = []
    
    driver = None
    try:
        for test in test_cases:
            # Setup driver for each test if none exists
            if not driver:
                website_url = test.get('website_url')
                driver = setup_driver(website_url, browser_type, headless)
            
            # Run the test
            result = run_test_case(driver, test)
            test['actual_result'] = result
            results.append(test)
    
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Error running tests: {str(e)}',
            'results': results
        }), 500
    
    finally:
        if driver:
            driver.quit()
    
    # Return results as JSON
    return jsonify({
        'success': True,
        'message': f'Successfully ran {len(results)} tests',
        'results': results
    })


if __name__ == "__main__":
    # Create test cases file if it doesn't exist
    if not os.path.exists(test_cases_file):
        pd.DataFrame(columns=["id", "description", "website_url", "username", "password", 
                              "username_field", "password_field", "submit_button", 
                              "success_indicators", "logout_url", "expected", 
                              "browser", "actual_result"]).to_csv(test_cases_file, dashboard=False)
    
    # Run the app
    app.run(debug=True)